# TO CREATE AND REBUILD VENV ON WINDOWS
# cd FlaskAPI
# virtualenv venv
# .\venv\Scripts\activate
# pip install -r requirements.txt

# TO ADD REQUIREMENTS
# .\venv\Scripts\activate
# pip install package_name
# pip freeze > requirements.txt


import uuid

from flask import Flask, request, jsonify, session
import os
import json
import secrets
import openai
from flask_cors import CORS
import uuid
from docx import Document
import firebase_admin
from firebase_admin import credentials, storage
import io
from datetime import datetime


app = Flask(__name__)

# Initialize Firebase Admin using environment variable or file
firebase_credentials_env = os.environ.get('FIREBASE_CREDENTIALS_JSON')
if firebase_credentials_env:
    cred = credentials.Certificate(json.loads(firebase_credentials_env))
else:
    cred = credentials.Certificate('FlaskAPI/firebasecred.json')

firebase_admin.initialize_app(cred, {'storageBucket': 'traitor-14f52.appspot.com'})
app.config['SESSION_COOKIE_NAME'] = 'session'
app.config['SESSION_TYPE'] = 'filesystem'
app.secret_key = secrets.token_hex(16)
# CORS(app, supports_credentials=True, origins=["DOMAIN"]) PROD MODE | Replace domain with versel or real domain
CORS(app, supports_credentials=True, origins="*")
api_key = os.environ.get('OPENAI_API_KEY', '')
openai.api_key = api_key


@app.before_request
def ensure_session_token():
    print("Before request triggered.")
    if 'sessionToken' not in session:
        print("Token not in session.")
        session['sessionToken'] = str(uuid.uuid4())
        print(f"New sessionToken generated: {session['sessionToken']}")
    else:
        print("Token found in session.")
        print(f"Existing sessionToken found: {session['sessionToken']}")



@app.route('/')
def hello_world():
    return 'chatgpt nice brother wojak!'

@app.route('/set-token')
def set_token():
    session['sessionToken'] = str(uuid.uuid4())
    return f"Token set: {session['sessionToken']}"

@app.route('/get-token')
def get_token():
    return f"{session.get('sessionToken', 'No token found.')}"

@app.route('/file-uploaded', methods=['POST'])
# USE /documentscan INSTEAD
def file_uploaded():
    data = request.json
    file_name = data['fileName']
    session_id = data['sessionID']
    # PROCESS WORD DOC HERE
    return jsonify({'success': True})

@app.route('/askgpt', methods=['POST'])
def askGPT():
    data = request.json
    user_prompt = data.get('prompt', '')

    messages = [
        {"role": "system",
         "content": "You are trying to determine whether or not ChatGPT wrote the prompt being given to you."},
        {"role": "user", "content": user_prompt}
    ]

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=messages
        )

        return jsonify({
            'testName': 'AskGPT',
            'success': True,
            'response': response['choices'][0]['message']['content'].strip()
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })


@app.route('/reverseprompt', methods=['POST'])
def reversePrompt():
    data = request.json
    user_prompt = data.get('prompt', '')

    description_messages = [
        {"role": "user",
         "content": "this was generated by chatgpt. what do you think the prompt that the user asked chat gpt was to generate this?\n\n" + user_prompt}
    ]

    try:
        description_response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=description_messages
        )
        description = description_response['choices'][0]['message']['content'].strip()

        reverse_messages = [
            {"role": "user", "content": description}
        ]

        reverse_response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=reverse_messages
        )
        reversed_prompt = reverse_response['choices'][0]['message']['content'].strip()

        return jsonify({
            'testName': 'ReversePrompt',
            'success': True,
            'original_prompt': user_prompt,
            'reversed_prompt': reversed_prompt,
            'description': description
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/documentscan', methods=['POST'])
def document_scan():
    data = request.json
    session_token = data.get('session_token')
    file_name = data.get('file_name')

    firebase_file_path = f"files/{session_token}/{file_name}"

    bucket = storage.bucket()
    blob = bucket.blob(firebase_file_path)
    file_blob = blob.download_as_bytes()

    file_stream = io.BytesIO(file_blob)
    metadata, full_text, analysis_result = extract_metadata_and_text(file_stream)

    return jsonify({"metadata": metadata, "text": full_text, "analysis": analysis_result})

def extract_metadata_and_text(file_stream):
    doc = Document(file_stream)

    core_properties = doc.core_properties
    attributes = [
        'title', 'author', 'created', 'modified', 'last_modified_by',
        'description', 'category', 'comments', 'subject', 'keywords',
        'version', 'revision', 'identifier', 'language', 'content_status'
    ]

    metadata = {}
    for attr in attributes:
        if hasattr(core_properties, attr):
            value = getattr(core_properties, attr)
            if value:
                metadata[attr] = str(value)

    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)

    analysis_result = analyze_metadata_with_chatgpt(metadata)

    return metadata, full_text, analysis_result

def analyze_metadata_with_chatgpt(metadata):
    prompt = f"Analyze the following Word document metadata for any odd or suspicious characteristics that may indicate cheating:\n{json.dumps(metadata, indent=2)}"

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=[{"role": "system", "content": prompt}]
        )
        analysis_result = response['choices'][0]['message']['content'].strip()
        return analysis_result

    except Exception as e:
        return f"Error in analyzing metadata: {str(e)}"

@app.route('/extract-text', methods=['POST'])
def extract_text():
    data = request.json
    session_token = data.get('session_token')
    file_name = data.get('file_name')

    firebase_file_path = f"files/{session_token}/{file_name}"

    bucket = storage.bucket()
    blob = bucket.blob(firebase_file_path)
    file_blob = blob.download_as_bytes()

    file_stream = io.BytesIO(file_blob)
    doc = Document(file_stream)

    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)

    return jsonify({"text": full_text})



if __name__ == '__main__':
    app.run(debug=True)
